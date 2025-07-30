from fpdf import FPDF
from datetime import datetime
import pandas as pd
import io
import xlsxwriter
from docx import Document
from docx.shared import Pt


class PDF(FPDF):
    def __init__(self, company_name, period_desc):
        super().__init__()
        self.company_name = company_name
        self.period_desc = period_desc

    def header(self):
        self.set_font('Helvetica', 'B', 12)
        self.cell(0, 10, self.company_name, 0, 1, 'L')
        self.set_font('Helvetica', '', 10)
        self.cell(0, 10, f"Variance Analysis Report: {self.period_desc}", 0, 1, 'L')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'L')
        self.cell(0, 10, f'Generated on: {datetime.now().strftime("%Y-%m-%d")}', 0, 0, 'R')


def generate_pdf_report(df, company_name, period_desc, dollar_thresh, percent_thresh):
    pdf = PDF(company_name, period_desc)
    pdf.alias_nb_pages()
    pdf.add_page()

    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Executive Summary', 0, 1)
    pdf.set_font('Helvetica', '', 10)
    summary_text = (
        f"This report details {len(df)} material financial variances for {company_name} "
        f"for the period of {period_desc}. Materiality was defined as a variance "
        f"exceeding ${dollar_thresh:,.0f} or {percent_thresh:.1f}%."
    )
    pdf.multi_cell(0, 5, summary_text)
    pdf.ln(10)

    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 10, 'Detailed Analysis', 0, 1)
    pdf.set_font('Helvetica', '', 10)

    # --- THIS IS THE CORRECTED LAYOUT LOGIC ---
    for index, row in df.iterrows():
        # Account Name and Variance Type
        pdf.set_font('Helvetica', 'B', 10)
        pdf.multi_cell(0, 5, f"Account: {row['Account Name']} ({row['Variance Type']})", ln=True)

        # Financial Details
        pdf.set_font('Helvetica', '', 9)
        details = (
            f"Current: ${row['Current Period']:,.2f}   |   "
            f"Prior: ${row['Prior Period']:,.2f}   |   "
            f"Change: ${row['Dollar Variance']:,.2f} ({row['Percent Variance']:.1f}%)"
        )
        pdf.multi_cell(0, 5, details, ln=True)

        # Explanation
        pdf.set_font('Helvetica', 'I', 9) # Italic for the explanation
        pdf.multi_cell(0, 5, f"Explanation: {row['Explanation']}", ln=True)

        # Add a space between entries
        pdf.ln(5)

    pdf_buffer = io.BytesIO()
    pdf.output(pdf_buffer)
    return pdf_buffer.getvalue()


def generate_excel_report(input_df, analysis_df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        input_df.to_excel(writer, sheet_name='Raw Data', index=False)
        analysis_df.to_excel(writer, sheet_name='Variance Analysis', index=False, startrow=2)

        workbook = writer.book
        worksheet = writer.sheets['Variance Analysis']

        header_format = workbook.add_format({'bold': True, 'bg_color': '#DDEBF7', 'border': 1})
        currency_format = workbook.add_format({'num_format': '$#,##0.00'})
        percent_format = workbook.add_format({'num_format': '0.00%'})
        wrap_format = workbook.add_format({'valign': 'top', 'text_wrap': True})

        for col_num, value in enumerate(analysis_df.columns.values):
            worksheet.write(2, col_num, value, header_format)

        worksheet.set_column('A:A', 30, wrap_format)  # Account
        worksheet.set_column('B:D', 15, currency_format)  # Financials
        worksheet.set_column('E:E', 12, percent_format)  # Percent
        worksheet.set_column('F:G', 12, wrap_format)  # Type, Material
        worksheet.set_column('H:H', 60, wrap_format)  # Explanation

    return output.getvalue()



def generate_word_report(df, company_name, period_desc):
    document = Document()
    document.add_heading(f"{company_name} - Variance Analysis Report", level=0)
    document.add_paragraph(f"Period: {period_desc} | Generated: {datetime.now().strftime('%Y-%m-%d')}")

    document.add_heading('Detailed Variance Analysis', level=1)

    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Account'
    hdr_cells[1].text = 'Variance'
    hdr_cells[2].text = 'Figures (Current | Prior)'
    hdr_cells[3].text = 'Explanation'

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Account Name'])
        row_cells[1].text = f"${row['Dollar Variance']:,.2f} ({row['Percent Variance']:.1f}%)"
        row_cells[2].text = f"${row['Current Period']:,.2f} | ${row['Prior Period']:,.2f}"
        row_cells[3].text = str(row['Explanation'])

    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()